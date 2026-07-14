#!/usr/bin/env python3
"""Build the frozen narrative-proposal reference DOCX for reader v2 formats."""

from datetime import datetime, timezone
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "editions/reader_manuscript/v2_0/profiles/reader-v2-reference.docx"
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)


def set_font(style, name: str, size: float, color: RGBColor | None = None) -> None:
    style.font.name = name
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = color


def spacing(style, before: float, after: float, line: float) -> None:
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = line


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    set_font(normal, "Calibri", 11)
    spacing(normal, 0, 8, 1.333)

    title = doc.styles["Title"]
    set_font(title, "Calibri", 30, DARK_BLUE)
    title.font.bold = True
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    spacing(title, 96, 8, 1.0)

    subtitle = doc.styles["Subtitle"]
    set_font(subtitle, "Calibri", 15, BLUE)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    spacing(subtitle, 0, 28, 1.0)

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[name]
        set_font(style, "Calibri", size, color)
        style.font.bold = True
        spacing(style, before, after, 1.0)
        style.paragraph_format.keep_with_next = True

    for name in ["List Bullet", "List Number"]:
        style = doc.styles[name]
        set_font(style, "Calibri", 11)
        spacing(style, 0, 4, 1.208)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)

    caption = doc.styles["Caption"]
    set_font(caption, "Calibri", 9, DARK_BLUE)
    caption.font.italic = True
    spacing(caption, 4, 6, 1.0)

    doc.core_properties.title = "The ASI Stack Reader v2 DOCX Reference"
    doc.core_properties.author = "Corben Sorenson"
    doc.core_properties.subject = "Frozen narrative-proposal style reference"
    frozen_time = datetime(2026, 7, 13, tzinfo=timezone.utc)
    doc.core_properties.created = frozen_time
    doc.core_properties.modified = frozen_time
    doc.add_paragraph("The ASI Stack", style="Title")
    doc.add_paragraph("Reader Edition v2.0", style="Subtitle")
    doc.add_heading("Reference heading", level=1)
    doc.add_paragraph("Reference body paragraph.")
    doc.add_paragraph("Reference bullet.", style="List Bullet")
    doc.add_paragraph("Reference numbered item.", style="List Number")
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    normalized = OUTPUT.with_suffix(".normalized.docx")
    with ZipFile(OUTPUT) as source, ZipFile(normalized, "w", ZIP_DEFLATED, compresslevel=9) as target:
        for name in source.namelist():
            info = ZipInfo(name, (1980, 1, 1, 0, 0, 0))
            info.compress_type = ZIP_DEFLATED
            info.external_attr = 0o644 << 16
            target.writestr(info, source.read(name))
    normalized.replace(OUTPUT)
    print(f"Reader v2 DOCX reference wrote: {OUTPUT}")


if __name__ == "__main__":
    build()
